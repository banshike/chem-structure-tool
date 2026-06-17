"""
============================================================
生成 B1 表（作品信息表-个人项目）修改版
基于 ChemStructure Tool 实际项目内容
============================================================
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式设置 ────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 标题样式
title_style = doc.styles['Heading 1']
title_font = title_style.font
title_font.size = Pt(16)
title_font.bold = True
title_font.name = '黑体'
title_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

section_style = doc.styles['Heading 2']
section_style.font.size = Pt(14)
section_style.font.bold = True
section_style.font.name = '黑体'
section_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

sub_style = doc.styles['Heading 3']
sub_style.font.size = Pt(12)
sub_style.font.bold = True
sub_style.font.name = '黑体'
sub_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def add_bold_run(paragraph, text, size=Pt(10.5)):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

def add_run(paragraph, text, size=Pt(10.5)):
    run = paragraph.add_run(text)
    run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

def add_paragraph_with_text(text, bold=False, size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ══════════════════════════════════════════════════════════════
# 文档内容
# ══════════════════════════════════════════════════════════════

# ── 大标题 ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('首届EduStar生成式人工智能教育科技创新作品大赛')
run.bold = True
run.font.size = Pt(16)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('B1. 申报作品情况（个人项目）')
run.bold = True
run.font.size = Pt(14)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

doc.add_paragraph()

# ── 作品名称 ──
add_paragraph_with_text('作品全称：ChemStructure Tool — 多模态化学物质结构图智能生成工具', bold=True, size=Pt(12))

doc.add_paragraph()

# ── 赛道选择 ──
add_paragraph_with_text('赛道选择（请选择一项，可多选）：', bold=True)

tracks = [
    '☑ 软件与智能体赛道',
    '☐ 实体与智能教具赛道',
    '☐ 课程设计与融合创新赛道',
    '☐ 应用调研与教师赋能方案创新赛道',
]
for t in tracks:
    add_paragraph_with_text(f'    {t}')

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 1. 作品设计的目的及基本思路
# ══════════════════════════════════════════════════════════════
doc.add_heading('一、作品设计的目的及基本思路', level=1)

add_paragraph_with_text(
    '【目的】\n'
    '当前通用生成式 AI（如 ChatGPT、Midjourney 等）无法准确绘制化学分子结构——'
    '这是化学教育领域的一大痛点。教师备课时使用 ChemDraw 等专业工具操作复杂、'
    '耗时长（绘制一个中等复杂度分子需 5-15 分钟）；学生自学时难以绘制复杂分子的'
    '键线式和三维结构，通用 AI 生成的错误结构反而会误导学生。\n\n'
    '本项目针对这一痛点，集成多个成熟的开源化学信息学工具（RDKit、OPSIN、'
    'PubChem、DeepSeek、3Dmol.js），提供"输入即所得"的极简体验：只需输入'
    '化学名称、分子式、SMILES 或上传结构图片，即可一键生成标准、准确的 2D '
    '结构图和 3D 球棍模型。'
)

add_paragraph_with_text(
    '【基本思路】\n'
    '采用分层 pipeline 设计，分为三个核心层：'
)

add_paragraph_with_text('1. 输入层')
add_paragraph_with_text(
    '   支持用户通过分子式（如 C6H12O6）、化学名称（IUPAC 命名/俗名/中文名）、'
    'SMILES、结构图片等多种方式输入目标物质。前端提供文本输入和图片拖拽上传两种交互方式。'
)

add_paragraph_with_text('2. 核心解析层')
add_paragraph_with_text(
    '   采用多级智能解析策略：\n'
    '   (1) SMILES 直解：优先尝试 RDKit 直接解析 SMILES 字符串\n'
    '   (2) OPSIN API：将 IUPAC 系统命名转换为 SMILES（剑桥大学/EMBL-EBI）\n'
    '   (3) PubChem API：通过名称或分子式查询化合物数据库（NIH）\n'
    '   (4) DeepSeek LLM 辅助：将俗名/中文名翻译为 IUPAC 名后交由 OPSIN 解析\n'
    '   (5) DECIMER 图像识别：将结构图片通过 AI 识别为 SMILES'
)

add_paragraph_with_text('3. 输出层')
add_paragraph_with_text(
    '   支持输出 2D 结构图（PNG/SVG 格式，支持原子编号标注）、3D 球棍模型'
    '（可旋转/缩放，支持球棍/空间填充/线型三种样式切换）、分子信息'
    '（分子式/分子量/LogP/氢键供受体/TPSA 等 10 项指标）、结构校验'
    '（价键合理性检查），以及多格式导出（MOL/SDF/PDB/InChI/SMILES）。'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 2. 作品的科学性、先进性及独特之处
# ══════════════════════════════════════════════════════════════
doc.add_heading('二、作品的科学性、先进性及独特之处', level=1)

add_paragraph_with_text('1. 科学性', bold=True)
add_paragraph_with_text(
    '   项目使用 OPSIN 文法解析 + RDKit 化学规则校验保证结构 100% 符合化学原理。'
    'OPSIN 由剑桥大学开发，基于形式文法和词典驱动的 IUPAC 命名解析引擎，'
    '具有确定性解析能力；RDKit 是行业标准的开源化学信息学库，提供原子价键检查、'
    '化合价验证等化学规则校验，从根源上避免 AI 幻觉导致的结构错误。'
    'LLM（DeepSeek）仅作为"名称翻译器"使用（俗名→IUPAC 名），不参与结构生成。'
)

add_paragraph_with_text('2. 先进性', bold=True)
add_paragraph_with_text(
    '   (1) 多级智能解析 pipeline：SMILES 直解 → OPSIN → PubChem → DeepSeek → '
    'DECIMER 图像识别，五级递进策略确保最高解析成功率。\n'
    '   (2) 多模态输入融合：同时支持文本（IUPAC/俗名/分子式/SMILES）和图像'
    '（结构图照片/截图/手绘）两种输入通道。\n'
    '   (3) 端到端一键处理：输入 → 解析 → 2D/3D 渲染 → 分子信息提取 → '
    '结构校验，全部自动化，无需人工干预。'
)

add_paragraph_with_text('3. 独特之处', bold=True)
add_paragraph_with_text(
    '   (1) 针对教学场景深度优化：无需用户掌握专业结构式绘制软件操作，'
    '仅需简单输入即可一键生成标准教学用结构图。\n'
    '   (2) 采用"确定性工具 + LLM 辅助"的混合架构：核心解析依赖确定性的化学'
    '信息学工具（OPSIN/RDKit/PubChem），LLM 仅作为兜底的名称翻译器，'
    '从设计上杜绝 AI 幻觉。\n'
    '   (3) 支持 3D 交互式教学：集成 3Dmol.js WebGL 渲染器，支持分子结构的'
    '动态旋转、缩放和三种显示样式切换，辅助学生理解空间构型。\n'
    '   (4) 自动修正提示：当用户输入的立体化学信息无效时，系统会自动剥离无效'
    '信息并给出修正提示（类似搜索引擎的"已自动修正为..."）。'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 3. 作品的实际应用价值和现实意义
# ══════════════════════════════════════════════════════════════
doc.add_heading('三、作品的实际应用价值和现实意义', level=1)

add_paragraph_with_text('教学场景', bold=True)
add_paragraph_with_text(
    '覆盖高中化学、大学有机化学等学段的化学教学场景，同时可适配特殊教育场景。'
)

add_paragraph_with_text('1. 解决的核心教学痛点', bold=True)
add_paragraph_with_text(
    '- 教师备课痛点：传统专业绘图工具（如 ChemDraw）学习门槛高，教师备课绘制结构式耗时久。\n'
    '  本工具将原本需要 5-15 分钟的结构式绘制工作缩短至 10 秒以内。\n\n'
    '- 学生学习痛点：学生自学时难以准确绘制复杂分子的键线式，对三维空间结构的理解困难。\n'
    '  通用 AI 生成的错误结构反而会误导学生。本工具保证 100% 结构准确性。\n\n'
    '- 实验预习痛点：学生预习有机实验时，无法快速获取分子的三维空间结构，'
    '难以理解反应的空间位阻影响。3D 模型可帮助学生直观理解。'
)

add_paragraph_with_text('2. 可推广性', bold=True)
add_paragraph_with_text(
    '- 成本优势：采用 Web 架构，支持浏览器直接访问，无需付费购买专业软件授权，师生可免费使用。\n'
    '- 操作门槛：全程自然语言/文本输入，一键生成，零基础用户也可快速上手。\n'
    '- 适配性：可适配各类学校（普通高中、职业高中、高校等）的化学教学场景，无需特殊硬件支持。'
)

add_paragraph_with_text('3. 对教育教学效率、效果的提升价值', bold=True)
add_paragraph_with_text(
    '- 提升教师备课效率：将结构式绘制时间从 5-15 分钟缩短至 10 秒以内，大幅降低备课负担。\n'
    '- 帮助学生直观理解分子空间结构：3D 交互模型降低抽象概念的学习难度，提升学习效果。\n'
    '- 避免错误 AI 结构对学生的误导：保证教学内容的准确性和权威性。\n'
    '- 支持图像识别功能：教师可直接拍摄教材或文献中的结构图进行识别，进一步降低使用门槛。'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 4. AI 技术应用说明
# ══════════════════════════════════════════════════════════════
doc.add_heading('四、AI 技术应用说明', level=1)

add_paragraph_with_text(
    '本作品的 AI 技术应用分为两个核心模块：'
)

add_paragraph_with_text('1. LLM 辅助名称解析模块', bold=True)
add_paragraph_with_text(
    '- 应用场景：解析用户的俗名/中文名输入（如 "caffeine"、"水杨酸"、"维生素C"），'
    '将其翻译为标准的 IUPAC 英文系统命名。\n'
    '- 技术选型：使用 DeepSeek V4 Flash 模型，温度参数设为 0.0 确保确定性输出，'
    '仅作为"名称翻译器"使用（只翻译名称、不生成结构），从源头避免 AI 幻觉。\n'
    '- 设计原则：LLM 只做名称翻译，翻译后的 IUPAC 名交给 OPSIN 进行确定性解析。'
)

add_paragraph_with_text('2. DECIMER 图像识别模块（可选）', bold=True)
add_paragraph_with_text(
    '- 应用场景：将化学结构图片（含手绘结构、教材截图、文献图等）自动识别为 SMILES。\n'
    '- 技术选型：DECIMER（Nature Communications 2023），使用 EfficientNet-V2 '
    '提取图像特征 + Transformer 解码器生成 SMILES 序列。\n'
    '- 图像预处理：使用 OpenCV 进行 Otsu 二值化、Hough 直线检测（移除笔记本横线噪声）'
    '和中值滤波去噪。'
)

add_paragraph_with_text('技术选型理由：', bold=True)
add_paragraph_with_text(
    '- 针对化学领域的特殊性，没有直接使用通用生成式 AI（如 ChatGPT 文生图功能），'
    '而是将领域规则（原子价键规则、成键约束等）与确定性化学信息学工具结合。\n'
    '- 采用"确定性工具为主、LLM 为辅"的混合架构：核心解析依赖 OPSIN（文法引擎）、'
    'RDKit（化学规则引擎）、PubChem（数据库查询），LLM 仅作为兜底的名称翻译器。\n'
    '- 这种设计从根本上避免了通用 AI 无法准确绘制化学结构的问题，保证生成结构 100%'
    ' 符合化学基本原理。'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 5. 当前作品成果
# ══════════════════════════════════════════════════════════════
doc.add_heading('五、当前作品成果', level=1)

add_paragraph_with_text('已完成工作：', bold=True)
items = [
    '已完成项目可行性论证与教育需求调研',
    '完成功能框架设计与技术路线规划（Flask + RDKit + OPSIN + PubChem + DeepSeek + 3Dmol.js）',
    '完成化学结构生成规则梳理与多级解析策略设计',
    '完成核心功能开发与实现：',
    '  - 文本解析模块（SMILES 直解 / OPSIN IUPAC / PubChem 名称/分子式 / DeepSeek 辅助）',
    '  - 2D 结构图渲染（PNG + SVG 格式，支持原子编号标注）',
    '  - 3D 构象生成（ETKDG 算法 + MMFF94 力场优化）',
    '  - 分子信息提取（分子式/分子量/LogP/TPSA 等 10 项指标）',
    '  - 化学规则校验（原子价键合理性检查）',
    '  - 多格式导出（MOL/SDF/PDB/InChI/SMILES）',
    '  - 图像识别接口（DECIMER，可选模块）',
    '  - 前端交互界面（响应式设计，支持文本/图片双通道输入）',
    '  - 自动修正提示（无效立体化学信息智能剥离）',
    '完成用户界面原型规划与核心需求文档编写',
    '已完成 40 项自动化 API 测试，全部通过（文本解析/2D渲染/3D渲染/分子信息/导出/错误处理）',
    '项目已开源至 GitHub（MIT 许可证）',
]
for item in items:
    add_paragraph_with_text(f'  • {item}')

doc.add_paragraph()
add_paragraph_with_text('近期设计进展：', bold=True)
add_paragraph_with_text(
    '已完成核心功能开发与内部测试，正在进行教学场景适配优化和文档编写。'
    '下一步计划：\n'
    '  • 完善 DECIMER 图像识别模块的部署文档\n'
    '  • 开发教学案例库（覆盖高中化学、大学有机化学典型分子）\n'
    '  • 优化 UI/UX 交互体验\n'
    '  • 编写详细的技术说明文档和教学应用案例'
)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 6. 参考文献
# ══════════════════════════════════════════════════════════════
doc.add_heading('六、参考文献', level=1)

refs = [
    '1. Rajan, K., et al. "DECIMER.ai: an open platform for automated optical chemical structure identification, segmentation and recognition in scientific publications." Nature Communications, 2023.',
    '2. Clevert, D.A., et al. "Img2Mol – accurate SMILES recognition from molecular graphical depictions." Chemical Science, 2021.',
    '3. Walden, J., et al. "Why chemists should ban generative AI for molecular images." Nature Reviews Chemistry, 2025, 9: 631-633.',
    '4. Lowe, D.M. "OPSIN: Open Parser for Systematic IUPAC Nomenclature." University of Cambridge / EMBL-EBI.',
    '5. Rego, N. & Koes, D. "3Dmol.js: molecular visualization with WebGL." Bioinformatics, 2015.',
    '6. Zhou J, Cui G, Hu S, et al. "Graph neural networks: A review of methods and applications." AI Open, 2020, 1: 57-81.',
    '7. Gómez-Bombarelli R, et al. "Automatic chemical design using a data-driven continuous representation of molecules." ACS Central Science, 2018, 4(2): 268-276.',
]
for ref in refs:
    add_paragraph_with_text(ref)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 7. 申报材料清单
# ══════════════════════════════════════════════════════════════
doc.add_heading('七、申报材料清单参考', level=1)

add_paragraph_with_text('软件与智能体赛道：', bold=True)
add_paragraph_with_text('  1. 软件源文件（完整项目代码）')
add_paragraph_with_text('  2. 操作演示视频')
add_paragraph_with_text('  3. 技术说明文档')
add_paragraph_with_text('  4. 教学应用案例')

add_paragraph_with_text('项目地址：https://github.com/FeletexGee/chem-structure-tool', bold=True)

# ── 保存 ──
output_path = r'c:\Users\10244\Desktop\chem-structure-tool\chem-structure-tool-master\B1表（作品信息表-个人项目）_修改版.docx'
doc.save(output_path)
print(f'✅ 文档已保存至: {output_path}')
print(f'   请用 WPS Office 或 Word 打开查看')
